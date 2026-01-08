"""
Tender AI Platform - Document Text Extraction Service
Supports PDF (digital/scanned), DOCX, XLSX - all in-memory

WORKFLOW:
1. First-page scan of ALL files → classify to find Avis
2. Full extraction of Avis ONLY
3. Store Avis text → Run AI pipeline
"""

import io
import re
from typing import Dict, Optional, Tuple, List
from dataclasses import dataclass
from enum import Enum
from loguru import logger

# Document processing
import pypdf
from docx import Document as DocxDocument
import openpyxl
import pandas as pd


class DocumentType(str, Enum):
    AVIS = "AVIS"
    RC = "RC"
    CPS = "CPS"
    ANNEXE = "ANNEXE"
    UNKNOWN = "UNKNOWN"


class ExtractionMethod(str, Enum):
    DIGITAL = "DIGITAL"
    OCR = "OCR"


@dataclass
class FirstPageResult:
    """Result of first-page scan for classification"""
    filename: str
    first_page_text: str
    document_type: DocumentType
    is_scanned: bool
    mime_type: str
    file_size_bytes: int
    success: bool
    error: Optional[str] = None


@dataclass
class ExtractionResult:
    """Result of full text extraction from a document"""
    filename: str
    document_type: DocumentType
    text: str
    page_count: Optional[int]
    extraction_method: ExtractionMethod
    file_size_bytes: int
    mime_type: str
    success: bool
    error: Optional[str] = None


# Classification keywords for document type detection
# Priority order matters - AVIS is checked FIRST
CLASSIFICATION_KEYWORDS = {
    DocumentType.AVIS: [
        "avis de consultation",
        "avis d'appel d'offres", 
        "avis d'appel",
        "avis appel offres",
        "avis ao",
        "avis",  # Generic - filename must contain "avis"
    ],
    DocumentType.RC: [
        "règlement de consultation",
        "reglement de consultation",
        "règlement de la consultation",
        "reglement de la consultation",
    ],
    DocumentType.CPS: [
        "cahier des prescriptions spéciales",
        "cahier des prescriptions speciales",
        "cahier des clauses",
    ],
    DocumentType.ANNEXE: [
        "annexe",
        "additif",
        "avenant",
    ]
}

# Filename patterns for classification (regex patterns)
FILENAME_PATTERNS = {
    DocumentType.AVIS: [
        r'\bavis\b',           # "avis" as whole word
        r'\bavis[\s_-]',       # "avis " or "avis_" or "avis-"
        r'[\s_-]avis\b',       # " avis" or "_avis" or "-avis"
        r'avis[\s_-]*(ar|fr)', # "avis ar" or "avis fr" (Arabic/French)
    ],
    DocumentType.RC: [
        r'\brc\b',             # "rc" as whole word
        r'\brcdp\b',           # "rcdp"
        r'\brcdg\b',           # "rcdg"
    ],
    DocumentType.CPS: [
        r'\bcps\b',            # "cps" as whole word
        r'\bccaf\b',           # "ccaf"
        r'\bcctp\b',           # "cctp" (cahier des clauses techniques)
    ],
    DocumentType.ANNEXE: [
        r'\bannexe\b',
    ]
}


def classify_document(text: str, filename: str = "", use_ai: bool = False, is_scanned: bool = False) -> DocumentType:
    """
    Classify document type by scanning first-page content and filename.
    Priority: AVIS > RC > CPS > ANNEXE
    
    Args:
        text: First page text content
        filename: Document filename
        use_ai: Whether to use AI classification as fallback
        is_scanned: Whether document is scanned (limits text for AI)
    """
    text_lower = text.lower()
    filename_lower = filename.lower()
    
    # Extract just the file name without path
    base_filename = filename_lower.split('/')[-1].split('\\')[-1]
    
    # PRIORITY 1: Check filename patterns (most reliable for Avis detection)
    for doc_type in [DocumentType.AVIS, DocumentType.RC, DocumentType.CPS, DocumentType.ANNEXE]:
        if doc_type in FILENAME_PATTERNS:
            for pattern in FILENAME_PATTERNS[doc_type]:
                if re.search(pattern, base_filename, re.IGNORECASE):
                    # For AVIS, make sure it's not RC/CPS file with "avis" in name
                    if doc_type == DocumentType.AVIS:
                        # Exclude if filename clearly indicates RC or CPS
                        if re.search(r'\b(rc|cps|ccaf|rcdp|rcdg)\b', base_filename):
                            continue
                    return doc_type
    
    # PRIORITY 2: Check text content keywords
    for doc_type in [DocumentType.AVIS, DocumentType.RC, DocumentType.CPS, DocumentType.ANNEXE]:
        if doc_type in CLASSIFICATION_KEYWORDS:
            for keyword in CLASSIFICATION_KEYWORDS[doc_type]:
                if keyword in text_lower:
                    return doc_type
    
    # PRIORITY 3: Use AI classification if enabled and text available
    if use_ai and text and len(text.strip()) > 20:
        ai_result = classify_document_with_ai(text, filename, is_scanned)
        if ai_result != DocumentType.UNKNOWN:
            return ai_result
    
    return DocumentType.UNKNOWN


def classify_document_with_ai(text: str, filename: str = "", is_scanned: bool = False) -> DocumentType:
    """
    Use DeepSeek AI to classify document type.
    For scanned documents, limits to first 500 words.
    
    Args:
        text: Document text content
        filename: Document filename
        is_scanned: Whether document is scanned (OCR'd)
    
    Returns:
        DocumentType classification
    """
    try:
        from openai import OpenAI
        from app.core.config import settings
        
        if not settings.DEEPSEEK_API_KEY:
            logger.warning("DeepSeek API key not configured, skipping AI classification")
            return DocumentType.UNKNOWN
        
        # For scanned docs, use first 500 words only
        if is_scanned:
            words = text.split()[:500]
            text_to_analyze = " ".join(words)
        else:
            # For digital docs, use first 2000 chars
            text_to_analyze = text[:2000]
        
        client = OpenAI(
            api_key=settings.DEEPSEEK_API_KEY,
            base_url=settings.DEEPSEEK_BASE_URL
        )
        
        system_prompt = """You are a document classifier for Moroccan government tender documents.

Classify the document into ONE of these categories:
- AVIS: Avis de consultation, avis d'appel d'offres, notice of tender, announcement
- RC: Règlement de consultation, consultation rules
- CPS: Cahier des prescriptions spéciales, specifications document
- ANNEXE: Annexe, addendum, modification document
- UNKNOWN: Cannot determine

RULES:
1. If document mentions "avis de consultation" or "avis d'appel d'offres" → AVIS
2. If document is primarily about rules/procedures for bidders → RC
3. If document contains technical specifications or requirements → CPS
4. If document modifies another document → ANNEXE
5. Only return UNKNOWN if truly cannot classify

Respond with ONLY one word: AVIS, RC, CPS, ANNEXE, or UNKNOWN"""

        user_content = f"""Filename: {filename}

Document first page content:
{text_to_analyze}

Classification:"""

        logger.info(f"AI classifying document: {filename}")
        
        response = client.chat.completions.create(
            model=settings.DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ],
            max_tokens=10,
            temperature=0
        )
        
        result = response.choices[0].message.content.strip().upper()
        
        # Map response to DocumentType
        type_map = {
            "AVIS": DocumentType.AVIS,
            "RC": DocumentType.RC,
            "CPS": DocumentType.CPS,
            "ANNEXE": DocumentType.ANNEXE,
        }
        
        doc_type = type_map.get(result, DocumentType.UNKNOWN)
        logger.info(f"AI classified {filename} as: {doc_type.value}")
        return doc_type
        
    except Exception as e:
        logger.error(f"AI classification failed: {e}")
        return DocumentType.UNKNOWN


# ============================
# FIRST-PAGE EXTRACTION (Classification Phase)
# ============================

def _is_pdf_scanned(file_bytes: io.BytesIO) -> Tuple[bool, str]:
    """
    Check if PDF is scanned by attempting digital extraction of first page.
    Returns (is_scanned, first_page_text)
    """
    file_bytes.seek(0)
    try:
        reader = pypdf.PdfReader(file_bytes)
        if len(reader.pages) == 0:
            return True, ""
        
        first_page_text = reader.pages[0].extract_text() or ""
        
        # If text is too sparse (<100 chars), it's scanned
        is_scanned = len(first_page_text.strip()) < 100
        return is_scanned, first_page_text
    except Exception as e:
        logger.warning(f"PDF scan check failed: {e}")
        return True, ""


def _ocr_first_page_pdf(file_bytes: io.BytesIO) -> str:
    """OCR only the first page of a scanned PDF"""
    try:
        from paddleocr import PaddleOCR
        import fitz  # PyMuPDF
        
        # Initialize PaddleOCR - only use_angle_cls and lang are universally supported
        ocr = PaddleOCR(use_angle_cls=True, lang='fr')
        
        file_bytes.seek(0)
        doc = fitz.open(stream=file_bytes.read(), filetype="pdf")
        
        if len(doc) == 0:
            doc.close()
            return ""
        
        # Only first page
        page = doc[0]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_bytes = pix.tobytes("png")
        
        result = ocr.ocr(img_bytes, cls=True)
        doc.close()
        
        if result and result[0]:
            return "\n".join([line[1][0] for line in result[0]])
        return ""
        
    except ImportError as e:
        logger.warning(f"OCR dependencies not available: {e}")
        return "[OCR REQUIRED]"
    except Exception as e:
        logger.error(f"First-page OCR failed: {e}")
        return ""


def _get_first_page_docx(file_bytes: io.BytesIO) -> str:
    """Get first ~1000 chars from DOCX (approximates first page)"""
    file_bytes.seek(0)
    try:
        doc = DocxDocument(file_bytes)
        text_parts = []
        char_count = 0
        
        for para in doc.paragraphs:
            text_parts.append(para.text)
            char_count += len(para.text)
            if char_count > 1000:
                break
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX first-page extraction failed: {e}")
        return ""


def _get_first_page_doc(file_bytes: io.BytesIO) -> str:
    """
    Extract first page text from legacy .doc files.
    Uses multiple fallback methods.
    """
    file_bytes.seek(0)
    content = file_bytes.read()
    
    # Method 1: Try using antiword via subprocess (if installed)
    try:
        import subprocess
        import tempfile
        
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            result = subprocess.run(
                ['antiword', tmp_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                import os
                os.unlink(tmp_path)
                # Return first 1000 chars
                return result.stdout[:1000]
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        finally:
            try:
                import os
                os.unlink(tmp_path)
            except:
                pass
    except Exception:
        pass
    
    # Method 2: Basic binary text extraction (fallback)
    try:
        # .doc files often have readable text mixed with binary
        text_parts = []
        # Try to decode as various encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                decoded = content.decode(encoding, errors='ignore')
                # Extract readable text sequences (4+ chars)
                import re
                words = re.findall(r'[a-zA-ZÀ-ÿ\s]{4,}', decoded)
                if words:
                    text = ' '.join(words)
                    # Clean up
                    text = re.sub(r'\s+', ' ', text).strip()
                    if len(text) > 50:  # Reasonable amount of text
                        return text[:1000]
            except:
                continue
    except Exception as e:
        logger.warning(f"Binary .doc extraction failed: {e}")
    
    return ""


def _extract_full_doc(file_bytes: io.BytesIO) -> Tuple[str, int]:
    """Full extraction from legacy .doc files"""
    file_bytes.seek(0)
    content = file_bytes.read()
    
    # Method 1: Try using antiword via subprocess
    try:
        import subprocess
        import tempfile
        
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            result = subprocess.run(
                ['antiword', tmp_path],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0 and result.stdout.strip():
                import os
                os.unlink(tmp_path)
                return result.stdout, None
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        finally:
            try:
                import os
                os.unlink(tmp_path)
            except:
                pass
    except Exception:
        pass
    
    # Method 2: Basic binary text extraction
    try:
        text_parts = []
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                decoded = content.decode(encoding, errors='ignore')
                import re
                words = re.findall(r'[a-zA-ZÀ-ÿ0-9\s\.,;:\-\(\)]{4,}', decoded)
                if words:
                    text = ' '.join(words)
                    text = re.sub(r'\s+', ' ', text).strip()
                    if len(text) > 100:
                        return text, None
            except:
                continue
    except Exception as e:
        logger.warning(f"Binary .doc full extraction failed: {e}")
    
    return "[.DOC EXTRACTION FAILED - Install antiword for better support]", None


def _get_first_page_xlsx(file_bytes: io.BytesIO) -> str:
    """Get first rows from first sheet of XLSX"""
    file_bytes.seek(0)
    try:
        wb = openpyxl.load_workbook(file_bytes, read_only=True, data_only=True)
        if not wb.sheetnames:
            wb.close()
            return ""
        
        sheet = wb[wb.sheetnames[0]]
        text_parts = []
        row_count = 0
        
        for row in sheet.iter_rows(values_only=True):
            row_values = [str(cell) if cell is not None else "" for cell in row]
            if any(row_values):
                text_parts.append(" | ".join(row_values))
                row_count += 1
                if row_count > 20:  # First 20 rows
                    break
        
        wb.close()
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"XLSX first-page extraction failed: {e}")
        return ""


def extract_first_page(filename: str, file_bytes: io.BytesIO, use_ai_classification: bool = True) -> FirstPageResult:
    """
    Extract FIRST PAGE ONLY for classification purposes.
    This is a quick scan to identify document type.
    
    Args:
        filename: Document filename
        file_bytes: Document content as BytesIO
        use_ai_classification: Whether to use AI for classification fallback
    """
    # Skip temp files and hidden files
    base_name = filename.split('/')[-1]
    if base_name.startswith('~$') or base_name.startswith('.'):
        return FirstPageResult(
            filename=filename,
            first_page_text="",
            document_type=DocumentType.UNKNOWN,
            is_scanned=False,
            mime_type="",
            file_size_bytes=0,
            success=False,
            error="Temporary or hidden file - skipped"
        )
    
    # Get file size
    file_bytes.seek(0, 2)
    file_size = file_bytes.tell()
    file_bytes.seek(0)
    
    # Determine MIME type from extension
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    mime_map = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'xls': 'application/vnd.ms-excel',
    }
    mime_type = mime_map.get(ext, 'application/octet-stream')
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    is_scanned = False
    first_page_text = ""
    
    try:
        if ext == 'pdf' or mime_type == 'application/pdf':
            is_scanned, first_page_text = _is_pdf_scanned(file_bytes)
            
            if is_scanned and not first_page_text:
                logger.info(f"Scanned PDF detected, OCR first page: {filename}")
                first_page_text = _ocr_first_page_pdf(file_bytes)
                
        elif ext == 'docx' or 'wordprocessingml' in mime_type:
            first_page_text = _get_first_page_docx(file_bytes)
            
        elif ext == 'doc':
            # Legacy .doc files - try extraction
            logger.info(f"Extracting legacy .doc file: {filename}")
            first_page_text = _get_first_page_doc(file_bytes)
            
        elif ext in ('xls', 'xlsx') or 'excel' in mime_type or 'spreadsheet' in mime_type:
            first_page_text = _get_first_page_xlsx(file_bytes)
            
        elif ext == 'txt' or mime_type == 'text/plain':
            file_bytes.seek(0)
            first_page_text = file_bytes.read(2000).decode('utf-8', errors='ignore')
            
        else:
            return FirstPageResult(
                filename=filename,
                first_page_text="",
                document_type=DocumentType.UNKNOWN,
                is_scanned=False,
                mime_type=mime_type,
                file_size_bytes=file_size,
                success=False,
                error=f"Unsupported file type: {ext}"
            )
        
        # Classify based on first-page content
        # Use AI classification for better accuracy, especially on scanned docs
        doc_type = classify_document(
            first_page_text, 
            filename, 
            use_ai=use_ai_classification,
            is_scanned=is_scanned
        )
        
        return FirstPageResult(
            filename=filename,
            first_page_text=first_page_text,
            document_type=doc_type,
            is_scanned=is_scanned,
            mime_type=mime_type,
            file_size_bytes=file_size,
            success=True
        )
        
    except Exception as e:
        logger.error(f"First-page extraction failed for {filename}: {e}")
        return FirstPageResult(
            filename=filename,
            first_page_text="",
            document_type=DocumentType.UNKNOWN,
            is_scanned=False,
            mime_type=mime_type,
            file_size_bytes=file_size,
            success=False,
            error=str(e)
        )


# ============================
# FULL EXTRACTION (Only for identified Avis)
# ============================

def _extract_full_pdf_digital(file_bytes: io.BytesIO) -> Tuple[str, int]:
    """Full digital extraction from PDF"""
    file_bytes.seek(0)
    reader = pypdf.PdfReader(file_bytes)
    page_count = len(reader.pages)
    
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        text_parts.append(page_text)
    
    return "\n\n".join(text_parts), page_count


def _extract_full_pdf_ocr(file_bytes: io.BytesIO) -> Tuple[str, int]:
    """Full OCR extraction from scanned PDF"""
    try:
        from paddleocr import PaddleOCR
        import fitz
        
        logger.info("Full OCR extraction starting...")
        
        # Initialize PaddleOCR - only use_angle_cls and lang are universally supported
        ocr = PaddleOCR(use_angle_cls=True, lang='fr')
        
        file_bytes.seek(0)
        doc = fitz.open(stream=file_bytes.read(), filetype="pdf")
        page_count = len(doc)
        
        all_text = []
        for page_num in range(page_count):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_bytes = pix.tobytes("png")
            
            result = ocr.ocr(img_bytes, cls=True)
            
            if result and result[0]:
                page_text = "\n".join([line[1][0] for line in result[0]])
                all_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        doc.close()
        return "\n\n".join(all_text), page_count
        
    except ImportError as e:
        logger.warning(f"OCR dependencies not available: {e}")
        return "[OCR REQUIRED - Dependencies not installed]", 0
    except Exception as e:
        logger.error(f"Full OCR failed: {e}")
        return f"[OCR FAILED: {str(e)}]", 0


def _extract_full_docx(file_bytes: io.BytesIO) -> Tuple[str, int]:
    """Full extraction from DOCX"""
    file_bytes.seek(0)
    doc = DocxDocument(file_bytes)
    
    paragraphs = [para.text for para in doc.paragraphs]
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            paragraphs.append(row_text)
    
    return "\n".join(paragraphs), None


def _extract_full_xlsx(file_bytes: io.BytesIO) -> Tuple[str, int]:
    """Full extraction from Excel"""
    file_bytes.seek(0)
    
    try:
        wb = openpyxl.load_workbook(file_bytes, read_only=True, data_only=True)
        
        all_text = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            all_text.append(f"=== Sheet: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(row_values):
                    all_text.append(" | ".join(row_values))
        
        wb.close()
        return "\n".join(all_text), None
        
    except Exception:
        file_bytes.seek(0)
        try:
            df = pd.read_excel(file_bytes, sheet_name=None)
            all_text = []
            for sheet_name, sheet_df in df.items():
                all_text.append(f"=== Sheet: {sheet_name} ===")
                all_text.append(sheet_df.to_string())
            return "\n".join(all_text), None
        except Exception as e:
            return f"[EXCEL EXTRACTION FAILED: {e}]", None


def extract_full_document(filename: str, file_bytes: io.BytesIO, is_scanned: bool = False) -> ExtractionResult:
    """
    Full extraction of a single document.
    Use appropriate method based on is_scanned flag.
    """
    file_bytes.seek(0, 2)
    file_size = file_bytes.tell()
    file_bytes.seek(0)
    
    mime_type = detect_mime_type(filename)
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    try:
        if ext == 'pdf' or mime_type == 'application/pdf':
            if is_scanned:
                text, page_count = _extract_full_pdf_ocr(file_bytes)
                method = ExtractionMethod.OCR
            else:
                text, page_count = _extract_full_pdf_digital(file_bytes)
                method = ExtractionMethod.DIGITAL
                
        elif ext == 'docx' or 'wordprocessingml' in mime_type:
            text, page_count = _extract_full_docx(file_bytes)
            method = ExtractionMethod.DIGITAL
            
        elif ext == 'doc':
            text, page_count = _extract_full_doc(file_bytes)
            method = ExtractionMethod.DIGITAL
            
        elif ext in ('xls', 'xlsx') or 'excel' in mime_type or 'spreadsheet' in mime_type:
            text, page_count = _extract_full_xlsx(file_bytes)
            method = ExtractionMethod.DIGITAL
            
        elif ext == 'txt' or mime_type == 'text/plain':
            file_bytes.seek(0)
            text = file_bytes.read().decode('utf-8', errors='ignore')
            page_count = None
            method = ExtractionMethod.DIGITAL
            
        else:
            return ExtractionResult(
                filename=filename,
                document_type=DocumentType.UNKNOWN,
                text="",
                page_count=None,
                extraction_method=ExtractionMethod.DIGITAL,
                file_size_bytes=file_size,
                mime_type=mime_type,
                success=False,
                error=f"Unsupported file type: {ext}"
            )
        
        doc_type = classify_document(text, filename)
        
        return ExtractionResult(
            filename=filename,
            document_type=doc_type,
            text=text,
            page_count=page_count,
            extraction_method=method,
            file_size_bytes=file_size,
            mime_type=mime_type,
            success=True
        )
        
    except Exception as e:
        logger.error(f"Full extraction failed for {filename}: {e}")
        return ExtractionResult(
            filename=filename,
            document_type=DocumentType.UNKNOWN,
            text="",
            page_count=None,
            extraction_method=ExtractionMethod.DIGITAL,
            file_size_bytes=file_size,
            mime_type=mime_type,
            success=False,
            error=str(e)
        )


# ============================
# MAIN WORKFLOW: Classify then Extract Avis
# ============================

def classify_all_documents(zip_files: Dict[str, io.BytesIO]) -> List[FirstPageResult]:
    """
    STEP 1: Scan first page of ALL files to classify them.
    Returns classification results (first-page text is temporary, will be discarded).
    """
    results = []
    
    for filename, file_bytes in zip_files.items():
        # Skip hidden files and directories
        if filename.startswith('.') or filename.startswith('__'):
            continue
        
        logger.info(f"Classifying: {filename}")
        result = extract_first_page(filename, file_bytes)
        results.append(result)
    
    return results


def find_avis_document(classifications: List[FirstPageResult]) -> Optional[FirstPageResult]:
    """
    STEP 2: Find the Avis de Consultation from classification results.
    Returns the Avis document info, or None if not found.
    """
    for result in classifications:
        if result.success and result.document_type == DocumentType.AVIS:
            logger.success(f"Found Avis document: {result.filename}")
            return result
    
    logger.warning("No Avis document found in ZIP")
    return None


def extract_avis_only(
    zip_files: Dict[str, io.BytesIO],
    avis_info: FirstPageResult
) -> Optional[ExtractionResult]:
    """
    STEP 3: Extract FULL content of Avis document only.
    Uses appropriate method (digital/OCR) based on classification.
    """
    if avis_info.filename not in zip_files:
        logger.error(f"Avis file not found in ZIP: {avis_info.filename}")
        return None
    
    file_bytes = zip_files[avis_info.filename]
    
    logger.info(f"Full extraction of Avis: {avis_info.filename} (scanned={avis_info.is_scanned})")
    return extract_full_document(avis_info.filename, file_bytes, avis_info.is_scanned)


def process_tender_zip(zip_files: Dict[str, io.BytesIO]) -> Tuple[Optional[ExtractionResult], List[FirstPageResult]]:
    """
    MAIN WORKFLOW: Process a tender ZIP file.
    
    1. Classify all documents (first-page scan)
    2. Find Avis de Consultation
    3. Extract full Avis content
    4. Return (avis_extraction, all_classifications)
    
    Classifications are returned for logging/debugging but their first_page_text
    should be discarded after processing.
    """
    # Step 1: Classify all documents
    logger.info("Phase 1: Classifying all documents...")
    classifications = classify_all_documents(zip_files)
    
    # Log classification results
    for c in classifications:
        status = "✓" if c.success else "✗"
        scanned = " [SCANNED]" if c.is_scanned else ""
        logger.info(f"  {status} {c.filename} → {c.document_type.value}{scanned}")
    
    # Step 2: Find Avis
    logger.info("Phase 2: Locating Avis document...")
    avis_info = find_avis_document(classifications)
    
    if not avis_info:
        return None, classifications
    
    # Step 3: Full extraction of Avis only
    logger.info("Phase 3: Full extraction of Avis...")
    avis_extraction = extract_avis_only(zip_files, avis_info)
    
    # Clear first-page texts from memory (they're no longer needed)
    for c in classifications:
        c.first_page_text = ""  # Discard
    
    return avis_extraction, classifications


# ============================
# LEGACY FUNCTION (for backward compatibility)
# ============================

def extract_all_from_zip(zip_files: Dict[str, io.BytesIO]) -> List[ExtractionResult]:
    """
    DEPRECATED: Old function that extracted all files fully.
    Kept for backward compatibility but should use process_tender_zip instead.
    """
    logger.warning("Using legacy extract_all_from_zip - consider using process_tender_zip")
    
    results = []
    for filename, file_bytes in zip_files.items():
        if filename.startswith('.') or filename.startswith('__'):
            continue
        
        # Check if scanned first
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        is_scanned = False
        if ext == 'pdf':
            is_scanned, _ = _is_pdf_scanned(file_bytes)
        
        result = extract_full_document(filename, file_bytes, is_scanned)
        results.append(result)
    
    return results
